from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PAPER = ROOT / "paper"
OUTPUT = PAPER / "PolyDegradeML_consolidated_manuscript.docx"

SOURCES = [
    ("PolyDegradeML Manuscript Draft", PAPER / "manuscript_draft.md", False),
    ("Appendix A. Evidence-Mapped Manuscript Outline", PAPER / "evidence_mapped_outline.md", True),
    ("Appendix B. Tables And Figures Plan", PAPER / "tables_and_figures_plan.md", True),
    ("Appendix C. Publication Readiness Checklist", PAPER / "publication_readiness_checklist.md", True),
    ("Appendix D. Missing Information To Ask Author", PAPER / "missing_information_to_ask_author.md", True),
    ("Appendix E. Paper Folder README", PAPER / "README.md", True),
    ("Appendix F. Original Paper Outline", PAPER / "outline.md", True),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=60, start=100, bottom=60, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_column_widths(table, widths) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = width
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
                tc_w.set(qn("w:type"), "dxa")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 2.0

    for name, size, before, after in [
        ("Heading 1", 16, 18, 8),
        ("Heading 2", 14, 14, 6),
        ("Heading 3", 12, 10, 4),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.2

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.5

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("PolyDegradeML consolidated manuscript draft")
    footer.runs[0].font.name = "Times New Roman"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(80, 80, 80)


def clean_heading(text: str) -> str:
    return text.strip().replace("`", "")


def add_formatted_runs(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\\*\\*[^*]+\\*\\*|\\*[^*]+\\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(10.5)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def parse_table(lines, start_idx):
    rows = []
    idx = start_idx
    while idx < len(lines) and lines[idx].lstrip().startswith("|"):
        raw = lines[idx].strip()
        cells = split_markdown_table_row(raw)
        rows.append(cells)
        idx += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        rows.pop(1)
    return rows, idx


def split_markdown_table_row(raw: str):
    text = raw.strip()
    if text.startswith("|"):
        text = text[1:]
    if text.endswith("|"):
        text = text[:-1]
    cells = []
    current = []
    in_code = False
    escaped = False
    for ch in text:
        if escaped:
            current.append(ch)
            escaped = False
            continue
        if ch == "\\":
            current.append(ch)
            escaped = True
            continue
        if ch == "`":
            in_code = not in_code
            current.append(ch)
            continue
        if ch == "|" and not in_code:
            cells.append("".join(current).strip())
            current = []
        else:
            current.append(ch)
    cells.append("".join(current).strip())
    return cells


def add_table(doc: Document, rows) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)

    if col_count == 4:
        widths = [Inches(1.35), Inches(2.15), Inches(2.0), Inches(1.0)]
    elif col_count == 3:
        widths = [Inches(1.7), Inches(3.1), Inches(1.7)]
    elif col_count == 2:
        widths = [Inches(2.2), Inches(4.3)]
    else:
        col_width = 6.5 / col_count
        widths = [Inches(col_width)] * col_count
    set_column_widths(table, widths)

    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[c_idx] if c_idx < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            add_formatted_runs(paragraph, text)
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(8.5)
                if r_idx == 0:
                    run.bold = True
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
    doc.add_paragraph()


def add_code_block(doc: Document, code_lines) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run("\n".join(code_lines))
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(9)


def add_markdown(doc: Document, path: Path, title: str, appendix: bool) -> None:
    if appendix:
        doc.add_section(WD_SECTION.NEW_PAGE)
        heading = doc.add_paragraph(style="Heading 1")
        add_formatted_runs(heading, title)
    else:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(12)
        title_para.paragraph_format.line_spacing = 1.15
        run = title_para.add_run(title)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(18)
        run.bold = True

    lines = path.read_text(encoding="utf-8").splitlines()
    idx = 0
    in_code = False
    code_lines = []

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        if not stripped:
            idx += 1
            continue

        if stripped.startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = clean_heading(heading_match.group(2))
            if not appendix and idx == 0:
                # The file's top heading duplicates the title page above.
                idx += 1
                continue
            style = "Heading 1" if level <= 2 else "Heading 2" if level == 3 else "Heading 3"
            p = doc.add_paragraph(style=style)
            add_formatted_runs(p, text)
            idx += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, re.sub(r"^[-*]\s+", "", stripped))
            idx += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            idx += 1
            continue

        p = doc.add_paragraph()
        add_formatted_runs(p, stripped)
        idx += 1


def main() -> None:
    doc = Document()
    configure_document(doc)

    for title, path, appendix in SOURCES:
        if path.exists():
            add_markdown(doc, path, title, appendix)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
